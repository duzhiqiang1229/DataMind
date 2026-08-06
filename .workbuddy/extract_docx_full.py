"""Extract all docx files to individual txt files for reading."""
import sys, os
from docx import Document

docs_dir = r"D:\DataMind\docs"
out_dir = r"D:\DataMind\.workbuddy\extracted"
os.makedirs(out_dir, exist_ok=True)

files = sorted([f for f in os.listdir(docs_dir) if f.endswith('.docx')])

for fname in files:
    fpath = os.path.join(docs_dir, fname)
    base_name = os.path.splitext(fname)[0]
    out_path = os.path.join(out_dir, base_name + ".txt")
    
    try:
        doc = Document(fpath)
        lines = []
        
        for para in doc.paragraphs:
            text = para.text
            if not text.strip():
                lines.append("")
                continue
            
            style_name = para.style.name if para.style else ""
            
            if "Title" in style_name or "title" in style_name:
                lines.append(f"# {text}")
            elif "Heading 1" in style_name:
                lines.append(f"# {text}")
            elif "Heading 2" in style_name:
                lines.append(f"## {text}")
            elif "Heading 3" in style_name:
                lines.append(f"### {text}")
            elif "Heading 4" in style_name:
                lines.append(f"#### {text}")
            elif "Heading 5" in style_name or "Heading 6" in style_name:
                lines.append(f"##### {text}")
            else:
                lines.append(text)
        
        # Tables
        for i, table in enumerate(doc.tables):
            lines.append(f"\n[表{i+1}]")
            for row in table.rows:
                cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
        
        with open(out_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))
        
        print(f"OK: {fname} -> {base_name}.txt ({len(lines)} lines)")
    except Exception as e:
        print(f"ERROR: {fname}: {e}")

print("\n=== ALL DONE ===")
