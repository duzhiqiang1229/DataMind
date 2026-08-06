import sys, os
from docx import Document

docs_dir = r"D:\DataMind\docs"
files = sorted([f for f in os.listdir(docs_dir) if f.endswith('.docx')])

for fname in files:
    fpath = os.path.join(docs_dir, fname)
    print(f"\n{'='*80}")
    print(f"FILE: {fname}")
    print(f"{'='*80}")
    try:
        doc = Document(fpath)
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                style = para.style.name if para.style else ""
                if "Heading" in style or "heading" in style or "Title" in style or "title" in style:
                    print(f"\n### {text}")
                else:
                    print(text)
        for i, table in enumerate(doc.tables):
            print(f"\n[TABLE {i+1}]")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                print(" | ".join(cells))
    except Exception as e:
        print(f"ERROR reading {fname}: {e}")

print("\n\n=== ALL FILES READ ===")
